"""FastAPI 主入口 — 论文审稿系统 API"""

from __future__ import annotations

import json
import os
import sys
import tempfile
import uuid
import asyncio
from contextlib import asynccontextmanager
from datetime import datetime, timezone, timedelta
from pathlib import Path

from dotenv import load_dotenv
from pathlib import Path
# .env 在项目根目录，不是 server/ 下
_env_path = Path(__file__).resolve().parent.parent / ".env"
_loaded = load_dotenv(dotenv_path=_env_path)  # 加载 .env 文件
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from openai import AsyncOpenAI
print(f"[Config] .env 文件: {_env_path}  加载={'成功' if _loaded else '失败'}", flush=True)

# ── 管理员认证配置 ─────────────────────────────────────
ADMIN_USERNAME = os.getenv("ADMIN_USERNAME", "admin")
ADMIN_PASSWORD = os.getenv("ADMIN_PASSWORD", "")
JWT_SECRET = os.getenv("JWT_SECRET", "")
JWT_EXPIRE_HOURS = int(os.getenv("JWT_EXPIRE_HOURS", "72"))

if not ADMIN_PASSWORD:
    print("[Security] ⚠️  请设置环境变量 ADMIN_USERNAME 和 ADMIN_PASSWORD 后再启动服务", flush=True)
    print("[Security] 例如: ADMIN_USERNAME=admin ADMIN_PASSWORD=your_strong_password", flush=True)
if not JWT_SECRET:
    print("[Security] ⚠️  JWT_SECRET 未设置，将使用默认值（生产环境强烈建议设置）", flush=True)
    print("[Security] 设置: JWT_SECRET=<your-256-bit-secret>", flush=True)

# 生成默认 JWT_SECRET
if not JWT_SECRET:
    import secrets
    JWT_SECRET = secrets.token_hex(32)
    print("[Security] 使用默认 JWT_SECRET（重启后失效）", flush=True)

# 依赖注入
security = HTTPBearer()

from fastapi import FastAPI, File, UploadFile, Form, HTTPException, Depends, status, Body
import pypdfium2 as pdfium

from models import (
    AIReviewItem,
    CompletionReport,
    CompletionItem,
    HistoryRecord,
    LogicalReview,
    CoherenceIssue,
    KnowledgeGraph,
    KnowledgeGraphNode,
    KnowledgeGraphEdge,
    ParsedDocument,
    ReviewSummary,
    Revision,
    RuleReport,
)
from ai_service.polisher import polish_paper as _polish_paper
from ai_service.lit_review import generate_literature_review as _gen_lit_review
from parser import parse_text, detect_missing_sections, extract_sections
from rule_engine.section_check import check_sections, analyze_section_balance
from rule_engine.format_check import check_format
from rule_engine.citation_check import check_citations

# 在全局状态区域添加
_models_cache: list[dict] | None = None  # 缓存模型列表
_models_cache_time: float = 0.0          # 缓存时间戳

# ── 已知模型列表（静态 + Ollama 发现）──────────────────────

_KNOWN_MODELS: list[str] = [
    "claude-sonnet-4-5-20251001",
    "claude-opus-4-5-20251001",
    "claude-haiku-4-5-20251001",
]

# 自定义模型端点（非默认 base_url 的模型）
# API key 优先从环境变量 LAN_MODEL_API_KEY 读取
_LAN_KEY = os.getenv("LAN_MODEL_API_KEY") or os.getenv("OPENAI_API_KEY") or "sk-litellmXa304304"
# 优先使用 LAN_MODEL_URL，回退到 OPENAI_BASE_URL，都未设置则用默认
_lan_url_raw = os.getenv("LAN_MODEL_URL") or os.getenv("OPENAI_BASE_URL") or "http://127.0.0.1:7000/v1"
_lan_url_base = _lan_url_raw.rstrip("/")
# 确保 URL 有 /v1 后缀（OpenAI 兼容 API 标准路径）
if not _lan_url_base.endswith("/v1"):
    _lan_url_base = _lan_url_base + "/v1"
_LAN_URL = _lan_url_base
_LAN_PROXY_ROOT = _LAN_URL.rsplit("/v1", 1)[0] if "/v1" in _LAN_URL else _LAN_URL
_CUSTOM_ENDPOINTS: dict[str, tuple[str, str]] = {
    "claude-sonnet-4-5-20251001": (_LAN_URL, _LAN_KEY),
    "claude-opus-4-5-20251001":   (_LAN_URL, _LAN_KEY),
    "claude-haiku-4-5-20251001":  (_LAN_URL, _LAN_KEY),
}
print(f"[Config] 代理端点 URL={_LAN_URL}  根路径={_LAN_PROXY_ROOT}  API Key={_LAN_KEY[:8]}...", flush=True)

def _discover_ollama_models() -> list[str]:
    """从 Ollama 发现可用模型，与已知模型列表合并"""
    import urllib.request
    try:
        resp = urllib.request.urlopen("http://localhost:11434/api/tags", timeout=2)
        data = resp.read().decode("utf-8")
        models = [m["name"] for m in json.loads(data).get("models", [])]
        seen: set[str] = set()
        result: list[str] = []
        for m in (models + _KNOWN_MODELS):
            if m not in seen:
                seen.add(m)
                result.append(m)
        return result
    except Exception:
        return list(_KNOWN_MODELS)


def _check_proxy_available(url: str, timeout: float = 3.0) -> bool:
    """检查代理端点是否可访问（优先使用 OpenAI 兼容的 /v1/models 端点）
    url 形如 http://server:7000/v1，需分别构造各探测路径。"""
    import urllib.request
    root = _LAN_PROXY_ROOT  # 不含 /v1 的根路径，如 http://server:7000
    api_key = _LAN_KEY if _LAN_KEY else "test-key"

    # 探测所有可能端点，返回第一个成功的
    probes = [
        (f"{root}/health", "GET", {}),
        (f"{root}/v1/health", "GET", {}),
        (f"{root}/v1/models", "GET", {"Authorization": f"Bearer {api_key}"}),
        (root, "GET", {}),
    ]
    for probe_url, method, headers in probes:
        try:
            req = urllib.request.Request(probe_url, method=method)
            for k, v in headers.items():
                req.add_header(k, v)
            resp = urllib.request.urlopen(req, timeout=timeout)
            print(f"[Models] 探测成功: {probe_url} → HTTP {resp.status}", flush=True)
            return True
        except Exception as e:
            print(f"[Models] 探测失败: {probe_url} → {type(e).__name__}: {e}", flush=True)
    return False


def _get_available_models(force_refresh: bool = False) -> list[dict]:
    """返回可用模型列表（带缓存，缓存有效期 5 分钟）。
    默认只做轻量探测：检查自定义代理端点是否可达，不可达则不返回该模型。
    """
    import time
    global _models_cache, _models_cache_time
    now = time.time()
    if not force_refresh and _models_cache is not None and (now - _models_cache_time) < 300:
        return _models_cache

    descriptions = {
        "claude-sonnet-4-5-20251001": "Claude Sonnet 4.5（局域网，深度推理）",
        "claude-opus-4-5-20251001": "Claude Opus 4.5（局域网，最强推理）",
        "claude-haiku-4-5-20251001": "Claude Haiku 4.5（局域网，极速响应）",
    }

    # 收集每个模型的代理端点
    model_endpoints: dict[str, str] = {}
    for model_name, (endpoint, _key) in _CUSTOM_ENDPOINTS.items():
        model_endpoints[model_name] = endpoint

    # 也合并 Ollama 本地模型
    ollama_models = _discover_ollama_models()

    result: list[dict] = []
    seen: set[str] = set()

    # 先检查自定义代理端点的可用性
    for model_name, endpoint in model_endpoints.items():
        if model_name in seen:
            continue
        seen.add(model_name)
        print(f"[Models] 正在探测代理 {endpoint} 上的模型 {model_name}...", flush=True)
        if _check_proxy_available(endpoint):
            result.append({"name": model_name, "desc": descriptions.get(model_name, f"模型 {model_name}")})
        else:
            print(f"[Models] 代理 {endpoint} 不可达，跳过模型 {model_name}", flush=True)

    # Ollama 本地模型直接包含（本地可用）
    for m in ollama_models:
        if m not in seen:
            seen.add(m)
            if m not in descriptions:
                result.append({"name": m, "desc": f"模型 {m}"})

    _models_cache = result
    _models_cache_time = now
    return result


def _create_llm(model_name: str | None = None) -> AsyncOpenAI:
    """根据模型名创建对应的 LLM 客户端（支持自定义端点）"""
    target_model = model_name or os.getenv("MODEL_NAME", "gpt-4o")

    # 自定义端点模型：使用独立 API 地址和 key
    if target_model in _CUSTOM_ENDPOINTS:
        endpoint, api_key = _CUSTOM_ENDPOINTS[target_model]
        client = AsyncOpenAI(api_key=api_key, base_url=endpoint)
        client._model = target_model  # type: ignore
        return client

    api_key = os.getenv("OPENAI_API_KEY", "NOT_SET")[:10]
    base_url = os.getenv("OPENAI_BASE_URL", "NOT_SET")
    print(f"[LLM] Creating client: key={api_key}..., base_url={base_url}", flush=True)
    client = AsyncOpenAI(
        api_key=os.getenv("OPENAI_API_KEY", ""),
        base_url=os.getenv("OPENAI_BASE_URL", "https://api.openai.com/v1"),
    )
    client._model = target_model  # type: ignore
    return client


# ── 全局状态 ───────────────────────────────────────────

_history: list[HistoryRecord] = []
_reports: dict[str, CompletionReport] = {}         # 存储完整报告供下载
_original_texts: dict[str, str] = {}               # 存储原始论文文本
_original_files: dict[str, tuple[bytes, str]] = {} # 存储原始文件字节和文件名（含扩展名）
_journals: dict[str, list[dict]] = {}              # 存储 LLM 推荐的期刊
_models_used: dict[str, str] = {}                   # 存储每个审稿使用的大模型名称

#PDF解析函数
def _extract_pdf_text(raw_bytes: bytes) -> str:
    """用 pypdfium2 从 PDF 提取纯文本（兼容 v3/v4/v5），并清理特殊字符"""
    import pypdfium2 as pdfium
    import re

    doc = pdfium.PdfDocument(raw_bytes)
    text_parts: list[str] = []
    for page in doc:
        page_text = ""
        try:
            textpage = page.get_textpage()
            n_chars = textpage.count_chars()
            if n_chars > 0:
                page_text = textpage.get_text_range(index=0, count=n_chars)
            textpage.close()
        except AttributeError:
            try:
                page_text = page.get_text()
            except AttributeError:
                page_text = page.get_textb().decode("utf-8", errors="replace")
        if page_text.strip():
            text_parts.append(page_text)
    doc.close()

    full_text = "\n".join(text_parts)

    # ── 🆕 PDF 文本预处理：移除可能导致问题的特殊字符 ──
    def clean_pdf_text(text: str) -> str:
        # 移除软连字符（PDF 断词标记）
        text = text.replace('\xad', '')
        # 移除零宽空格、零宽连接符等不可见字符
        text = re.sub(r'[\u200B-\u200F\u2028-\u202F\uFEFF]', '', text)
        # 移除所有控制字符（保留 \n \r \t）
        text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', text)
        # 压缩连续空行（保留最多两个）
        text = re.sub(r'\n{3,}', '\n\n', text)
        return text

    return clean_pdf_text(full_text)

#DOCX 解析函数
def _extract_docx_text(raw_bytes: bytes) -> str:
    """用 python-docx 提取文本"""
    from io import BytesIO
    from docx import Document as DocxDocument
    doc = DocxDocument(BytesIO(raw_bytes))
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

# ── DOCX 报告生成 ────────────────────────────────────
def _build_docx(report: CompletionReport, original_text: str, journals: list[dict] | None = None) -> bytes:
    """根据审稿报告、原始文本和 LLM 推荐的期刊生成带批注的 DOCX 文件"""
    from io import BytesIO
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    import re
    import traceback

    # ---------- 强化清洗函数 ----------
    def clean(text) -> str:
        """强力清洗，移除所有可能使 XML 崩溃的字符"""
        if text is None:
            return ""
        if not isinstance(text, str):
            text = str(text)
        # 1. 移除所有控制字符（除了 \n \r \t）
        text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', text)
        # 2. 移除未配对的 surrogate（会导致 XML 错误）
        text = re.sub(r'[\uD800-\uDFFF]', '', text)
        # 3. 移除零宽字符和不可见分隔符
        text = re.sub(r'[\u200B-\u200F\u2028-\u202F\uFEFF]', '', text)
        # 4. 移除其他可能导致问题的 Unicode 控制字符
        text = re.sub(r'[\uFFF0-\uFFFF]', '', text)
        # 5. 移除连续的过多换行（保留最多2个）
        text = re.sub(r'\n{3,}', '\n\n', text)
        return text

    # ---------- 安全写入辅助函数 ----------
    def safe_add_run(paragraph, text, **kwargs):
        """安全添加 run，失败时返回空 run"""
        if not text:
            return paragraph.add_run("", **kwargs)
        try:
            return paragraph.add_run(clean(text), **kwargs)
        except Exception as e:
            print(f"[DOCX] add_run 失败: {e}")
            return paragraph.add_run("", **kwargs)

    def safe_add_paragraph(doc, text, style=None):
        """安全添加段落，失败时返回空段落"""
        if not text:
            return doc.add_paragraph()
        try:
            return doc.add_paragraph(clean(text), style=style)
        except Exception as e:
            print(f"[DOCX] add_paragraph 失败: {e}")
            # 降级：尝试分段添加
            p = doc.add_paragraph()
            for line in str(text).split('\n')[:50]:  # 最多50行
                try:
                    if line.strip():
                        p.add_run(clean(line))
                        p.add_run('\n')
                except:
                    pass
            return p

    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    style = doc.styles['Normal']
    style.font.name = '等线'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5

    # ── 封面 ──────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('学术论文审稿报告')
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    safe_add_run(subtitle, f'论文：{report.file_name}')
    subtitle.runs[-1].font.size = Pt(14)
    subtitle.runs[-1].font.color.rgb = RGBColor(0x52, 0x52, 0x5B)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(clean(f'审稿时间：{report.timestamp.strftime("%Y-%m-%d %H:%M")}'))
    run.font.size = Pt(11)
    run = meta.add_run(clean(f'\n总评分数：{report.summary.overall_score:.0f} / 100'))
    run.font.size = Pt(11)

    rec_text = {
        "accept": "建议接收", "minor_revision": "小修后接收",
        "major_revision": "大修后复审", "reject": "不建议接收"
    }.get(report.summary.recommendation, report.summary.recommendation)
    run = meta.add_run(clean(f'\n审稿结论：{rec_text}'))
    run.font.size = Pt(11)

    doc.add_page_break()

    # ── 总体评价 ────────────────────────────────────
    doc.add_heading('一、总体评价', level=1)
    doc.add_heading('论文优点', level=2)
    for s in report.summary.strengths:
        safe_add_paragraph(doc, s, style='List Bullet')
    doc.add_heading('需要改进', level=2)
    for w in report.summary.weaknesses:
        safe_add_paragraph(doc, w, style='List Bullet')

    doc.add_page_break()

    # ── 逻辑连贯性审查 ────────────────────────────
    doc.add_heading('三、逻辑连贯性审查', level=1)
    if report.logical_review:
        if report.logical_review.research_theme:
            doc.add_heading('研究主题分析', level=2)
            safe_add_paragraph(doc, report.logical_review.research_theme)

        if report.logical_review.research_framework:
            doc.add_heading('研究路线图 / 整体框架', level=2)
            safe_add_paragraph(doc, report.logical_review.research_framework)

        if report.logical_review.overall_assessment:
            doc.add_heading('总体评价', level=2)
            safe_add_paragraph(doc, report.logical_review.overall_assessment)

        if report.logical_review.section_logic:
            doc.add_heading('章节与段落逻辑', level=2)
            for item in report.logical_review.section_logic:
                safe_add_paragraph(doc, item, style='List Bullet')

        if report.logical_review.argument_logic:
            doc.add_heading('论点与论据逻辑', level=2)
            for item in report.logical_review.argument_logic:
                safe_add_paragraph(doc, item, style='List Bullet')

        if report.logical_review.coherence_issues:
            doc.add_heading('检测到的逻辑问题', level=2)
            sev_label = {"error": "❌ 严重", "warning": "⚠️ 一般", "info": "ℹ️ 提示"}
            for i, ci in enumerate(report.logical_review.coherence_issues, 1):
                doc.add_heading(clean(f'{i}. {sev_label.get(ci.severity, ci.severity)} — {ci.issue_type}'), level=3)
                safe_add_paragraph(doc, f'位置：{ci.location}')
                safe_add_paragraph(doc, ci.description)
                if ci.suggestion:
                    p = doc.add_paragraph()
                    safe_add_run(p, f'💡 建议：{ci.suggestion}')
                    p.runs[-1].font.italic = True

        if report.logical_review.theme_consistency:
            doc.add_heading('主题一致性评价', level=2)
            for item in report.logical_review.theme_consistency:
                safe_add_paragraph(doc, item, style='List Bullet')

        if report.logical_review.knowledge_graph and (report.logical_review.knowledge_graph.nodes or report.logical_review.knowledge_graph.edges):
            doc.add_heading('知识图谱', level=2)
            if report.logical_review.knowledge_graph.summary:
                safe_add_paragraph(doc, report.logical_review.knowledge_graph.summary)

            doc.add_heading('节点列表', level=3)
            type_labels = {"theory": "理论", "method": "方法", "concept": "概念", "result": "结果", "variable": "变量", "finding": "发现"}
            for node in report.logical_review.knowledge_graph.nodes:
                label = f"{node.label} ({type_labels.get(node.type, node.type)})"
                if node.description:
                    safe_add_paragraph(doc, f"{label}: {node.description}", style='List Bullet')
                else:
                    safe_add_paragraph(doc, label, style='List Bullet')

            if report.logical_review.knowledge_graph.edges:
                doc.add_heading('关系列表', level=3)
                for edge in report.logical_review.knowledge_graph.edges:
                    safe_add_paragraph(doc, f"{edge.source} {edge.label} {edge.target}", style='List Bullet')
    else:
        safe_add_paragraph(doc, '暂无逻辑连贯性审查结果。')

    doc.add_page_break()

    # ── AI 审阅意见 ────────────────────────────────
    doc.add_heading('四、AI 审阅意见', level=1)
    if report.ai_reviews:
        for i, r in enumerate(report.ai_reviews, 1):
            doc.add_heading(clean(f'{i}. {r.section}'), level=2)
            safe_add_paragraph(doc, r.review_comment)
            if r.suggestion:
                p = doc.add_paragraph()
                safe_add_run(p, f'💡 {r.suggestion}')
                p.runs[-1].font.italic = True
                p.runs[-1].font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
    else:
        safe_add_paragraph(doc, '暂无 AI 审阅意见。')

    doc.add_page_break()

    # ── 修订痕迹 ────────────────────────────────────
    doc.add_heading('五、修订痕迹（修改前后对比）', level=1)
    if report.revisions:
        for i, r in enumerate(report.revisions, 1):
            rev_emoji = {"insertion": "➕ 新增", "deletion": "❌ 删除", "modification": "🔄 修改"}
            label = rev_emoji.get(r.revision_type.value, r.revision_type.value)
            doc.add_heading(clean(f'{i}. {label} — {r.location}'), level=2)

            if r.original_text:
                p = doc.add_paragraph()
                safe_add_run(p, '【原文】')
                p.runs[-1].font.bold = True
                p.runs[-1].font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
                p = doc.add_paragraph()
                safe_add_run(p, r.original_text)
                if p.runs:
                    p.runs[-1].font.strike = True
                    p.runs[-1].font.color.rgb = RGBColor(0xDC, 0x26, 0x26)

            p = doc.add_paragraph()
            safe_add_run(p, '【修改后】')
            p.runs[-1].font.bold = True
            p.runs[-1].font.color.rgb = RGBColor(0x05, 0x96, 0x69)
            p = doc.add_paragraph()
            safe_add_run(p, r.new_text)
            if p.runs:
                p.runs[-1].font.color.rgb = RGBColor(0x05, 0x96, 0x69)

            if r.rationale:
                p = doc.add_paragraph()
                safe_add_run(p, f'📝 理由：{r.rationale}')
                p.runs[-1].font.italic = True
                p.runs[-1].font.size = Pt(10)
    else:
        safe_add_paragraph(doc, '暂无修订痕迹。')

    doc.add_page_break()

    # ── 自动补全 ────────────────────────────────────
    doc.add_heading('六、自动补全内容', level=1)
    if report.completions:
        for i, c in enumerate(report.completions, 1):
            doc.add_heading(clean(f'{i}. {c.section}（置信度 {c.confidence:.0%}）'), level=2)
            safe_add_paragraph(doc, c.generated_content)
    else:
        safe_add_paragraph(doc, '论文章节完整，无需补全内容。')

    doc.add_page_break()

    # ── 推荐期刊 Top 10 ───────────────────────────────
    doc.add_heading('七、推荐期刊（Top 10）', level=1)

    score = report.summary.overall_score
    if score < 50:
        note = "当前稿件得分偏低，建议进一步完善实验验证和章节内容后投稿。"
    elif score < 70:
        note = "当前稿件处于中等水平，建议加强方法描述和实验对比后冲击高影响力期刊。"
    else:
        note = "稿件质量良好，建议优先尝试CCF-A/SCI Q1期刊，同时准备1-2个备选。"

    p = doc.add_paragraph()
    safe_add_run(p, f'💡 {note}')
    if p.runs:
        p.runs[-1].font.italic = True
        p.runs[-1].font.size = Pt(10)

    used_journals = (journals or [])[:10]
    if used_journals:
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Light Grid Accent 1'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        hdr = table.rows[0].cells
        hdr[0].text = '期刊名称'
        hdr[1].text = '级别'
        hdr[2].text = '匹配度'
        hdr[3].text = '推荐理由'

        for j in used_journals:
            row = table.add_row().cells
            row[0].text = clean(str(j.get("name", "")))
            p = row[1].paragraphs[0]
            safe_add_run(p, clean(str(j.get("level", ""))))
            if p.runs:
                p.runs[-1].font.bold = True
            match_val = j.get("match", "")
            row[2].text = clean(f"{match_val}%" if isinstance(match_val, (int, float)) else str(match_val))
            row[3].text = clean(str(j.get("reason", j.get("desc", ""))))

        from docx.shared import Cm
        for row_obj in table.rows:
            row_obj.cells[0].width = Cm(4.5)
            row_obj.cells[1].width = Cm(2.5)
            row_obj.cells[2].width = Cm(1.5)
            row_obj.cells[3].width = Cm(7.5)
    else:
        safe_add_paragraph(doc, '暂无推荐期刊数据。')

    doc.add_page_break()

    # ── 论文润色 ────────────────────────────────
    doc.add_heading('八、论文润色', level=1)
    if report.polished_paper:
        safe_add_paragraph(doc, report.polished_paper)
    else:
        safe_add_paragraph(doc, '暂无论文润色结果。')

    doc.add_page_break()

    # ── 文献综述 ────────────────────────────────
    doc.add_heading('九、文献综述', level=1)
    if report.literature_review:
        safe_add_paragraph(doc, report.literature_review)
    else:
        safe_add_paragraph(doc, '暂无文献综述结果。')

    # 写入内存
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()

def _extract_core_sections(text: str, sections: list, max_chars: int = 20000) -> str:
    """
    智能提取论文的核心章节，优先保留摘要、引言、方法、实验、结论。
    自动删除参考文献、致谢、附录等非核心内容。
    返回的文本长度不超过 max_chars。
    """
    # 章节优先级（数字越小越重要）
    priority_map = {
        "摘要": 1, "abstract": 1,
        "引言": 2, "introduction": 2, "绪论": 2,
        "背景": 2, "background": 2,
        "方法": 3, "methodology": 3, "methods": 3,
        "实验": 4, "experiment": 4,
        "结果": 4, "results": 4,
        "评估": 4, "evaluation": 4,
        "讨论": 5, "discussion": 5,
        "分析": 5, "analysis": 5,
        "结论": 1, "conclusion": 1, "总结": 1,
    }

    # 需要完全跳过的章节关键词
    skip_keywords = [
        "reference", "参考文献", "致谢", "acknowledgment",
        "appendix", "附录", "作者简介", "biography"
    ]

    # 收集并评分每个章节
    scored_sections = []
    for sec in sections:
        title_lower = sec.title.lower()
        if any(kw in title_lower for kw in skip_keywords):
            continue
        # 检查章节内容是否为空
        content = getattr(sec, 'content', '')
        if not content or len(content.strip()) < 10:
            continue
        # 计算优先级
        priority = 9
        for key, p in priority_map.items():
            if key in title_lower:
                priority = p
                break
        scored_sections.append((priority, sec))

    # 按优先级排序
    scored_sections.sort(key=lambda x: x[0])

    # 按优先级拼接，直到达到 max_chars
    result_parts = []
    total_len = 0
    for _, sec in scored_sections:
        content = getattr(sec, 'content', '')
        section_text = f"## {sec.title}\n{content}\n\n"
        if total_len + len(section_text) > max_chars:
            remaining = max_chars - total_len
            if remaining > 500:
                result_parts.append(section_text[:remaining])
            break
        result_parts.append(section_text)
        total_len += len(section_text)

    if not result_parts:
        return text[:max_chars]
    return "".join(result_parts)

# ── 语言检测 ──────────────────────────────────────────────────

def _is_english_text(text: str) -> bool:
    """判断是否为英文稿件（中文字符占比低于15%即为英文）"""
    import re
    chinese_count = len(re.findall(r'[一-鿿]', text))
    total = len(text.replace('\n', '').replace(' ', ''))
    if total == 0:
        return False
    return (chinese_count / total) < 0.15


def _build_ai_review_prompt(
    rule_lines: str, core_text: str, is_english: bool
) -> str:
    """构建 AI 审阅 prompt，英文稿件额外请求中文翻译"""

    bilingual_hint = ""
    if is_english:
        bilingual_hint = (
            "\n\n## 双语输出要求\n"
            "以下字段必须**同时提供英文原文和中文翻译**：logical_review 下所有字段（research_theme、research_framework、section_logic、argument_logic、coherence_issues.description、coherence_issues.suggestion、theme_consistency、overall_assessment、knowledge_graph.summary）。\n"
            "格式：在 JSON 字符串值内，用 `\\n\\n--- 中文翻译 ---\\n\\n`（两个字符）作为分隔符，不要使用字面量换行。\n"
            "示例 research_theme：\n"
            "\"research_theme\": \"[English analysis...]\\n\\n--- 中文翻译 ---\\n\\n[中文分析]\"\n"
            "示例 section_logic：\n"
            "\"section_logic\": [\"[English item 1...]\\n\\n--- 中文翻译 ---\\n\\n[中文条目 1]\"]\n"
            "注意：JSON 字符串中**绝对不要使用字面量换行符**（即不要按 Enter），所有换行必须写成 `\\n` 两个字符。\n"
            "以下字段**只用英文**（不需要翻译）：summary（strengths、weaknesses、overall_score、recommendation）、ai_reviews（section、review_comment、suggestion）、revisions、completions、journals、knowledge_graph.nodes、knowledge_graph.edges 的内部字段。\n"
            "请确保 JSON 格式合法，不要省略任何引号、逗号、花括号或方括号。\n"
            "请确保中文翻译准确、流畅，不遗漏关键信息。\n"
        )

    return f"""\
你是一位资深的学术期刊审稿人，擅长全面、深入地评审学术论文。请对以下论文进行审阅，返回详尽的审阅结果（纯 JSON，不要使用代码块）。

**重点权重分配（严格遵守）：**
- 🔬 **创新性/ Novelty（最高权重）**：重点评估论文的创新性，包括研究问题、方法、实验设计、结论等方面的原创贡献。
- 🧠 **逻辑性/ Logic（高权重）**：深度审查研究逻辑，包括论证链条、因果关系、推理严密性、结论合理性。
- 📐 章节/格式/段落长度（最低权重）：除非存在**严重**缺失或格式问题，否则**不要**在 weaknesses 或建议中提及章节不完整、段落过长、格式不规范等次要问题。

## 规则检查结果
{rule_lines}

## 论文核心章节（已自动筛选，保留摘要、引言、方法、实验、结论）

{core_text}

## 重要要求
你必须返回一个完整的 JSON 对象，包含以下全部字段。每项内容都必须**详尽、具体、有深度**，不能敷衍了事或泛泛而谈。

**JSON 格式要求：** 返回的必须是合法的 JSON。字符串中的换行必须使用 `\n`（两个字符：反斜杠+n），**不要使用字面量换行符**。双语字段中的 `--- 中文翻译 ---` 分隔符必须作为字符串的一部分包含在 JSON 值内（使用 `\n\n--- 中文翻译 ---\n\n` 格式）。

{bilingual_hint}
### summary - 总体评价
{{"overall_score": 0-100 整数, "strengths": ["优点1", "优点2", ...], "weaknesses": ["缺点1", "缺点2", ...], "recommendation": "accept|minor_revision|major_revision|reject"}}
要求：strengths 至少 4 条，weaknesses 至少 4 条，每条都要具体针对论文内容。**重点审查：创新性（研究问题的新颖性、方法的原创性、结论的贡献）和逻辑性（论证链条是否严密、因果关系是否成立、推理是否合理）。除非存在严重问题，不要在 weaknesses 中提及章节缺失、格式问题或段落长短。**

### ai_reviews - 逐章节 AI 审阅意见（至少 5 条）
每条格式：{{"section": "章节名", "review_comment": "详尽的审阅意见（≥200字）", "original_text": "原文关键片段（可选）", "suggestion": "具体的修改建议（≥50字）"}}
**审阅重点：创新性贡献和逻辑严密性，忽略格式和章节完整性等次要问题。**

### revisions - 修订痕迹（至少 8 条）
每条格式：{{"revision_type": "insertion|deletion|modification", "original_text": "原文（删除/修改时填写）", "new_text": "修改后的内容", "location": "位置描述", "rationale": "修改理由"}}
**修订重点：提升创新性表达和逻辑严密性，不要为格式或章节完整性提出修订建议。**

### completions - 自动补全内容（至少 2 条，如果章节完整则生成内容补充建议）
每条格式：{{"section": "章节名", "generated_content": "补充内容草稿（≥300字）", "confidence": 0.5-0.9}}
**补全重点：如果存在内容不足，优先补充能增强创新性论述和逻辑链条的内容。不要为缺少的格式性章节生成补全。**

### journals - 推荐期刊 Top 10
每条格式：{{"name": "期刊/会议全称", "level": "CCF-A/B/C 或 SCI Q1/Q2/Q3", "match": "匹配度百分比", "reason": "推荐理由（≥50字）"}}

### logical_review - 逻辑连贯性审查

**核心要求：** 以下所有审查必须紧密围绕论文的研究主题展开。请先仔细阅读全文，然后完成以下步骤：

**步骤1：提炼研究主题** — 概括论文的核心研究问题、研究目标、采用的主要方法。

**步骤2：绘制研究路线图 / 整体框架** — 这是审查的核心。必须根据论文的实际内容，用 ASCII 结构图清晰画出：
- 各章节/模块之间的逻辑关系（用箭头 → 表示流程方向）
- 研究的整体流程：问题提出 → 方法设计 → 实验/分析 → 结论
- 各模块之间的数据流或逻辑依赖关系

**ASCII 图示例格式：**
```
[研究背景与问题] → [文献综述] → [理论框架/假设提出]
                                      ↓
                              [研究方法设计] → [数据收集]
                                      ↓
                              [实验分析/模型构建] → [结果验证]
                                      ↓
                              [结论与建议]
```
请用论文中**实际的章节标题和内容**替换抽象的描述，不要只写通用模板。图中的每个节点应对应论文中的实际章节或模块。

**步骤3：基于研究主题和框架，逐项审查。**

### logical_review - 逻辑连贯性审查（注意：每个字段值内必须同时包含英文和中文翻译）

{{
  "research_theme": "[English analysis of the paper's research theme, core question, objective and methodology, ≥100 words]\\n\\n--- 中文翻译 ---\\n\\n[中文分析，至少100字]",
  "research_framework": "[English description of research roadmap and framework with ASCII diagram showing actual chapter names and logical flows]\\n\\n--- 中文翻译 ---\\n\\n[中文研究路线图与框架描述]",
  "section_logic": ["[English item 1 about section logic...]\\n\\n--- 中文翻译 ---\\n\\n[中文条目1]", "[English item 2...]\\n\\n--- 中文翻译 ---\\n\\n[中文条目2]", "[English item 3...]\\n\\n--- 中文翻译 ---\\n\\n[中文条目3]"],
  "argument_logic": ["[English item 1 about argument logic...]\\n\\n--- 中文翻译 ---\\n\\n[中文条目1]", "[English item 2...]\\n\\n--- 中文翻译 ---\\n\\n[中文条目2]", "[English item 3...]\\n\\n--- 中文翻译 ---\\n\\n[中文条目3]"],
  "coherence_issues": [
    {{"location": "位置描述（英文）", "issue_type": "section_logic|argument_logic|sentence_coherence|theme_mismatch", "description": "[English problem description...]\\n\\n--- 中文翻译 ---\\n\\n[中文问题描述]", "severity": "error|warning|info", "suggestion": "[English suggestion...]\\n\\n--- 中文翻译 ---\\n\\n[中文建议]"}}
  ],
  "theme_consistency": ["[English item 1 about theme consistency...]\\n\\n--- 中文翻译 ---\\n\\n[中文条目1]", "[English item 2...]\\n\\n--- 中文翻译 ---\\n\\n[中文条目2]"],
  "overall_assessment": "[English overall assessment, ≥150 words]\\n\\n--- 中文翻译 ---\\n\\n[中文总体评价，至少150字]",

  "knowledge_graph": {{
    "summary": "[English summary of the knowledge structure]\\n\\n--- 中文翻译 ---\\n\\n[中文知识图谱结构描述]",
    "nodes": [
      {{"id": "n1", "label": "[Entity name]", "type": "theory|method|concept|result|variable|finding", "description": "[Brief description]"}},
      {{"id": "n2", "label": "[Entity name]", "type": "theory|method|concept|result|variable|finding", "description": "[Brief description]"}},
      {{... 继续提取 8-15 个关键知识节点}}
    ],
    "edges": [
      {{"source": "n1", "target": "n2", "label": "[Relationship description]", "type": "supports|uses|contradicts|related|causes|improves"}},
      {{... 提取节点之间的关系，5-10 条}}
    ]
  }}
}}

注意：
- nodes: 提取论文中的关键知识元素，包括理论、方法、概念、结果、变量、发现等
- edges: 描述节点之间的关系，如"方法A支持理论B"、"变量C影响结果D"
- id 必须全局唯一，使用 "n1", "n2" 等简单命名
- edges 的 source 和 target 必须引用存在的节点 id
"""


# ── 新增：AI 审阅独立函数（抽离 prompt 逻辑）────────────────
async def _perform_ai_review(
    text: str,
    rules: list[RuleReport],
    llm: AsyncOpenAI,
    model: str,
    parsed: ParsedDocument,
    is_english: bool = False,
) -> dict:
    """
    仅执行 AI 审阅，返回解析后的 JSON dict。
    失败时返回兜底提取的数据，而不是抛出异常。
    """
    rule_lines = "\n".join(f"- [{r.severity.value}] {r.title}" for r in rules) or "无规则问题"
    section_titles = [s.title for s in parsed.sections] if parsed.sections else ["（未识别到章节）"]
    core_text = _extract_core_sections(text, parsed.sections, max_chars=20000)

    prompt = _build_ai_review_prompt(rule_lines, core_text, is_english)

    raw = "{}"
    try:
        print("[AI review] 开始执行AI审阅")
        model_to_use = getattr(llm, "_model", model or "gpt-4o")
        import asyncio

        # 🆕 根据模型动态设置超时时间
        if "claude" in model_to_use.lower():
            timeout_seconds = 1200.0  # Claude 需要更多时间（20 分钟）
        else:
            timeout_seconds = 900.0  # 其他模型（15 分钟）

        print(f"[AI review] 开始调用 LLM（超时 {timeout_seconds} 秒）...")
        try:
            resp = await asyncio.wait_for(
                llm.chat.completions.create(
                    model=model_to_use,
                    messages=[{"role": "user", "content": prompt}],
                    temperature=0.7,
                    max_tokens=32768,
                ),
                timeout=timeout_seconds
            )
            print("[AI review] LLM 响应已返回")
        except asyncio.TimeoutError:
            print(f"[AI review] ❌ 请求超时（{timeout_seconds} 秒）")
            raise RuntimeError(f"AI审阅超时（{timeout_seconds}秒），请检查模型服务是否正常")

        raw = resp.choices[0].message.content or "{}"
        import re as _re

        # 清洗思维链
        _think_end = raw.find('</think>')
        if _think_end >= 0:
            raw = raw[_think_end + len('</think>'):].strip()
        _think_match = _re.search(r'(Here\'s a thinking process:).*?\n\n', raw, _re.DOTALL | _re.IGNORECASE)
        if _think_match:
            raw = raw[_think_match.end():].strip()

        # 提取 JSON
        json_str = raw
        m = _re.search(r'```(?:json)?\s*\n(.*?)\n```', raw, _re.DOTALL)
        if m:
            json_str = m.group(1)
        else:
            _brace_start = raw.find('{')
            _brace_end = raw.rfind('}')
            if _brace_start >= 0 and _brace_end > _brace_start:
                json_str = raw[_brace_start:_brace_end + 1]

        # ---- 清洗函数 ----
        def clean_json(s: str) -> str:
            s = _re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', s)
            s = _re.sub(r',\s*}', '}', s)
            s = _re.sub(r',\s*]', ']', s)
            s = _re.sub(r'}\s*{', '},{', s)
            s = _re.sub(r'\]\s*\[', '],[', s)
            open_braces = s.count('{') - s.count('}')
            open_brackets = s.count('[') - s.count(']')
            if open_braces > 0:
                s += '}' * open_braces
            if open_brackets > 0:
                s += ']' * open_brackets
            return s

        json_str = clean_json(json_str)

        # ---- 兜底提取函数 ----
        def extract_from_raw_text(text: str) -> dict:
            """当 JSON 完全无法解析时，从原始文本中提取关键信息"""
            import re

            # 提取分数
            score = 60.0
            score_patterns = [
                r'(?:overall_score|整体评分|总分)[\s:]+(\d+\.?\d*)',
                r'分数[：:]\s*(\d+\.?\d*)',
                r'(\d+)分',
                r'评分[：:]\s*(\d+)',
            ]
            for pattern in score_patterns:
                match = re.search(pattern, text, re.IGNORECASE)
                if match:
                    score = float(match.group(1))
                    if score < 10:
                        score = score * 10
                    break
            if score < 0 or score > 100:
                score = 60.0

            # 提取优点
            strengths = []
            strength_patterns = [
                r'(?:strengths|优点|优势)[\s:]*\[(.*?)\]',
                r'优点[：:]\s*(.+?)(?=缺点|不足|局限)',
                r'1\.\s*(.+?)(?=2\.|\n)',
            ]
            for pattern in strength_patterns:
                match = re.search(pattern, text, re.DOTALL | re.IGNORECASE)
                if match:
                    items = re.split(r'[,，]\s*', match.group(1).strip())
                    strengths = [s.strip() for s in items if s.strip()]
                    break
            if not strengths:
                strengths = ["稿件结构完整", "实验设计合理"]

            # 提取缺点
            weaknesses = []
            weakness_patterns = [
                r'(?:weaknesses|缺点|不足|局限)[\s:]*\[(.*?)\]',
                r'缺点[：:]\s*(.+?)(?=优点|优势|建议)',
                r'2\.\s*(.+?)(?=3\.|\n)',
            ]
            for pattern in weakness_patterns:
                match = re.search(pattern, text, re.DOTALL | re.IGNORECASE)
                if match:
                    items = re.split(r'[,，]\s*', match.group(1).strip())
                    weaknesses = [s.strip() for s in items if s.strip()]
                    break
            if not weaknesses:
                weaknesses = ["建议补充相关工作章节", "建议完善实验对比分析"]

            if score >= 80:
                recommendation = "accept"
            elif score >= 70:
                recommendation = "minor_revision"
            else:
                recommendation = "major_revision"

            return {
                "summary": {
                    "overall_score": score,
                    "strengths": strengths[:5],
                    "weaknesses": weaknesses[:5],
                    "recommendation": recommendation,
                },
                "ai_reviews": [{
                    "section": "整体（AI审阅摘要）",
                    "review_comment": text[:4000] if text else "无法解析审阅结果",
                    "original_text": None,
                    "suggestion": "由于 JSON 解析失败，以上是从原始响应中提取的内容摘要。建议重新提交审稿或更换模型。",
                }],
                "revisions": [],
                "completions": [],
                "logical_review": None,
            }

        # ---- 尝试解析 ----
        parsed_json = None

        # 尝试1：直接解析
        try:
            parsed_json = json.loads(json_str)
            print("[AI review] JSON 解析成功（完整）")
        except json.JSONDecodeError as e:
            print(f"[AI review] 完整JSON解析失败: {e}")

        # 尝试2：strict=False
        if parsed_json is None:
            try:
                parsed_json = json.loads(json_str, strict=False)
                print("[AI review] JSON 解析成功（strict=False）")
            except json.JSONDecodeError as e:
                print(f"[AI review] strict=False 解析失败: {e}")

        # 尝试3：截断长度
        if parsed_json is None:
            for max_len in [len(json_str) // 2, len(json_str) // 4, len(json_str) // 8]:
                try:
                    parsed_json = json.loads(json_str[:max_len], strict=False)
                    print(f"[AI review] JSON 解析成功（截断到 {max_len}）")
                    break
                except json.JSONDecodeError:
                    continue

        # 尝试4：ast.literal_eval
        if parsed_json is None:
            try:
                import ast
                parsed_json = ast.literal_eval(json_str)
                print("[AI review] ast.literal_eval 解析成功")
            except Exception as e:
                print(f"[AI review] ast.literal_eval 也失败: {e}")

        # 尝试5：替换单引号
        if parsed_json is None:
            try:
                json_str_fixed = json_str.replace("'", '"')
                parsed_json = json.loads(json_str_fixed, strict=False)
                print("[AI review] 单引号替换后解析成功")
            except Exception as e:
                print(f"[AI review] 单引号替换后解析也失败: {e}")

        # ---- 所有尝试都失败 → 兜底提取 ----
        if parsed_json is None:
            print("[AI review] ⚠️ 所有解析均失败，使用兜底提取")
            parsed_json = extract_from_raw_text(json_str)

        return parsed_json

    except Exception as e:
        # 最外层异常捕获
        print(f"[AI review] 发生异常，使用兜底提取: {e}")
        try:
            return extract_from_raw_text(raw)
        except Exception:
            return {
                "summary": {
                    "overall_score": 60.0,
                    "strengths": ["稿件结构完整"],
                    "weaknesses": ["AI审阅解析失败，建议人工复审"],
                    "recommendation": "major_revision",
                },
                "ai_reviews": [{
                    "section": "整体",
                    "review_comment": "AI 审阅响应未能解析，请检查模型服务或重新提交。",
                    "suggestion": "建议检查模型 API 是否正常工作，或更换其他模型重新审稿。",
                }],
                "revisions": [],
                "completions": [],
                "logical_review": None,
            }


# ── 新增：降级兜底 JSON 构建函数 ──────────────────────────

def _build_fallback_review(text: str, parsed: ParsedDocument, rules: list[RuleReport]) -> dict:
    """当 AI 审阅失败时，生成默认的 parsed_json"""
    missing = detect_missing_sections(text)
    section_titles = [s.title for s in parsed.sections] if parsed.sections else []
    topic_clues = "、".join(section_titles[:6]) if section_titles else "未识别"

    parsed_json = {
        "summary": {
            "overall_score": 50.0 if missing else 70.0,
            "strengths": ["稿件结构基本完整", f"包含 {len(section_titles)} 个章节"],
            "weaknesses": [f"缺少以下章节：{', '.join(missing)}"] if missing else ["暂无需要改进之处"],
            "recommendation": "major_revision" if missing else "minor_revision",
        },
        "ai_reviews": [{
            "section": "整体",
            "review_comment": (
                f"【内容质量】\n稿件共 {parsed.word_count} 字，{len(parsed.sections)} 个章节"
                f"（{topic_clues}）。"
                + (f"缺失章节：{', '.join(missing)}。" if missing else "")
                + "\n\n【写作水平】\n文本整体可读，但建议进一步优化表达和逻辑连贯性。"
                "\n\n【具体问题】\n由于 LLM 审阅未能成功完成，建议人工补充审阅以下方面："
                "内容的创新性、实验设计的合理性、数据分析的准确性。"
                "\n\n【优点亮点】\n论文整体框架完整，章节划分合理。"
            ),
            "suggestion": (
                f"1. 补充缺失章节（{', '.join(missing) if missing else '无' }）\n"
                "2. 优化各章节之间的逻辑衔接\n"
                "3. 检查实验数据的完整性和准确性\n"
                "4. 完善文献综述，突出研究创新点"
            ),
        }],
        "revisions": [],
        "completions": [{
            "section": m,
            "generated_content": (
                f"【章节定位】\n「{m}」是学术论文中的核心组成部分，"
                f"对于完整呈现研究工作和说服审稿人具有关键作用。"
                f"当前稿件中缺少该章节，需要基于论文已有内容（主题：{topic_clues}）进行补充。\n\n"
                f"【核心要点】\n该章节应涵盖以下关键内容：\n"
                f"1. 与该章主题直接相关的背景阐述\n"
                f"2. 结合论文已有研究内容的深入分析\n"
                f"3. 支撑结论的关键论据和数据\n\n"
                f"【草稿正文】\n（请作者根据论文具体内容进行撰写，"
                f"建议围绕{topic_clues}等已有章节中涉及的相关内容展开，"
                f"确保与全文逻辑一致、风格统一。）"
            ),
            "confidence": 0.35,
        } for m in missing],
    }
    return parsed_json


# ── 新增：安全包装润色和综述 ──────────────────────────────
async def _safe_polish(
    text: str, sections: list, llm: AsyncOpenAI, model: str, is_english: bool = False
) -> str | None:
    """安全执行论文润色，失败返回 None"""
    try:
        print("[Polisher] 开始执行润色")
        import re

        # 1. 清洗不可见字符
        text = text.replace('\xad', '')  # 软连字符
        text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F]', '', text)  # 控制字符
        text = re.sub(r'\n{3,}', '\n\n', text).strip()

        # 2. 使用全文 + 章节列表，由 polish_paper 按章节分别润色后拼接
        print(f"[Polisher] 全文长度: {len(text)} 字，章节数: {len(sections)}")
        return await _polish_paper(text, sections, llm, model, is_english)
    except Exception as e:
        print(f"[Polisher] 失败: {e}", file=sys.stderr)
        return None

async def _safe_lit_review(
    text: str, llm: AsyncOpenAI, model: str, parsed: ParsedDocument, is_english: bool = False
) -> str | None:
    """安全执行文献综述，失败返回 None"""
    try:
        print("[LitReview] 开始执行文献综述")

        # 【关键】只提取核心章节，压缩到 10000 字
        core_text = _extract_core_sections(text, parsed.sections, max_chars=10000)

        # 如果提取失败（sections 为空），fallback 到截断
        if not core_text or len(core_text) < 500:
            core_text = text[:12000]
            print(f"[LitReview] 章节提取失败，使用截断: {len(core_text)} 字")
        else:
            print(f"[LitReview] 输入文本从 {len(text)} 字压缩到 {len(core_text)} 字")

        return await _gen_lit_review(core_text, llm, model, is_english)
    except Exception as e:
        print(f"[LitReview] 失败: {e}", file=sys.stderr)
        return None

# ── 审阅流程（核心改造）───────────────────────────────────

async def run_review(parsed: ParsedDocument, model_name: str | None = None, review_id: str | None = None) -> CompletionReport:
    """完整的审稿流程：规则检查 → 并行执行（AI审阅、润色、综述）→ 组装报告"""
    # 使用传入的 review_id 或生成新的
    _review_id = review_id or str(uuid.uuid4())[:8]
    text = parsed.full_text

    # 1. 规则检查（同步，轻量）
    rules: list[RuleReport] = []
    rules.extend(check_sections(text))
    rules.extend(analyze_section_balance([s.model_dump() for s in parsed.sections]))
    rules.extend(check_format(text))
    rules.extend(check_citations(text))

    # ── 语言检测 ──────────────────────────────────────────
    is_english = _is_english_text(text)
    if is_english:
        print("[Review] 检测到英文稿件，将生成中英文双语审阅结果")
    else:
        print("[Review] 检测到中文稿件，生成中文审阅结果")

    # 准备 LLM 客户端
    llm = _create_llm(model_name)
    model_to_use = getattr(llm, "_model", model_name or "gpt-4o")

    # LiteLLM 代理不支持并发，改为串行执行
    import time
    sem = asyncio.Semaphore(1)

    async def _limited_task(coro):  # 不变
        async with sem:
            return await coro

    # 调用方式完全一样，只是函数内部做了优化
    ai_task = _limited_task(_perform_ai_review(text, rules, llm, model_to_use, parsed, is_english))
    polish_task = _limited_task(_safe_polish(text, parsed.sections, llm, model_to_use, is_english))
    lit_task = _limited_task(_safe_lit_review(text, llm, model_to_use, parsed, is_english))

    task_start = time.time()
    print(f"[Timer] 三个并行任务开始发起: {task_start:.2f}")

    print("[Timer] 三个任务已全部发起，等待完成...")

    # 并行执行（完全不变）
    ai_result, polished_text, lit_review_text = await asyncio.gather(
        ai_task, polish_task, lit_task,
        return_exceptions=True
    )

    # ── 检测 LLM 失败 ──────────────────────────────
    llm_success = True
    error_messages: list[str] = []

    if isinstance(ai_result, Exception):
        llm_success = False
        error_messages.append(f"AI 审阅失败: {ai_result}")
    if isinstance(polished_text, Exception):
        llm_success = False
        error_messages.append(f"论文润色失败: {polished_text}")
    if isinstance(lit_review_text, Exception):
        llm_success = False
        error_messages.append(f"文献综述生成失败: {lit_review_text}")

    if not llm_success:
        print(f"[Review] ⚠️ LLM 审稿失败: {error_messages}")
        from fastapi import JSONResponse
        return JSONResponse(
            status_code=502,
            content={
                "error": "LLM 审稿失败，无法生成审稿报告",
                "messages": error_messages,
            },
        )

    # 处理各任务结果
    ai_error = ai_result if isinstance(ai_result, Exception) else None
    polish_error = polished_text if isinstance(polished_text, Exception) else None
    lit_error = lit_review_text if isinstance(lit_review_text, Exception) else None

    # 处理 AI 审阅结果
    if ai_error is not None:
        print(f"[AI Review] 失败，使用降级数据: {ai_error}", file=sys.stderr)
        parsed_json = _build_fallback_review(text, parsed, rules)
    else:
        parsed_json = ai_result

    if polished_text is None:
        polished_text = None  # 明确标记为 None
    if lit_review_text is None:
        lit_review_text = None

    # ── 降级补充：如果 LLM 未返回足够的 revisions/completions/journals ──
    # 复用原有逻辑（从 rules 生成修订，从缺失章节生成补全，从引擎推荐期刊）
    def _gen_revisions_from_rules() -> list[dict]:
        revs = []
        for r in rules:
            if r.severity in ("error", "warning"):
                revs.append({
                    "revision_type": "modification",
                    "original_text": r.description,
                    "new_text": r.suggestion or f"请根据「{r.title}」的要求修改此处内容。",
                    "location": r.location or r.category.value,
                    "rationale": f"规则检查 [{r.severity.value}] {r.title}",
                })
        return revs

    def _gen_completions_from_missing() -> list[dict]:
        missing = detect_missing_sections(text)
        completions = []
        for m in missing:
            completions.append({
                "section": m,
                "generated_content": (
                    f"【章节定位】\n"
                    f"「{m}」是学术论文中的关键组成部分，对于完整呈现研究工作、"
                    f"说服审稿人和读者具有重要作用。该章节在提交稿件中缺失，需要补充。\n\n"
                    f"【核心要点】\n"
                    f"该章节应涵盖以下内容：\n"
                    f"1. 与「{m}」相关的核心概念和背景知识\n"
                    f"2. 紧密结合论文已有研究内容的深入分析和讨论\n"
                    f"3. 支撑论文结论的关键论据、数据或引用\n\n"
                    f"【草稿正文】\n"
                    f"（请作者根据论文的具体研究内容、方法和数据进行撰写，"
                    f"确保与全文在逻辑上和风格上保持一致。）"
                ),
                "confidence": 0.30,
            })
        return completions

    from journal_recommender import recommend_journals as _recommend_journals

    def _gen_journals_from_content(text: str, summary: dict) -> list[dict]:
        score = summary.get("overall_score", 60)
        return _recommend_journals(text, overall_score=float(score))

    ai_revisions = parsed_json.get("revisions", []) or []
    ai_completions = parsed_json.get("completions", []) or []
    ai_journals = parsed_json.get("journals", []) or []
    sm = parsed_json.get("summary", {})

    if len(ai_revisions) < 2:
        rule_revs = _gen_revisions_from_rules()
        existing_locs = {r.get("location", "") for r in ai_revisions if isinstance(r, dict)}
        for rr in rule_revs:
            if rr.get("location") not in existing_locs:
                ai_revisions.append(rr)
        parsed_json["revisions"] = ai_revisions

    if len(ai_completions) < 1:
        rule_completions = _gen_completions_from_missing()
        existing_secs = {c.get("section", "") for c in ai_completions if isinstance(c, dict)}
        for rc in rule_completions:
            if rc.get("section") not in existing_secs:
                ai_completions.append(rc)
        parsed_json["completions"] = ai_completions

    if len(ai_journals) < 10:
        fallback_journals = _gen_journals_from_content(text, sm)
        existing_names = {j.get("name", "") for j in ai_journals if isinstance(j, dict)}
        for fj in fallback_journals:
            if fj.get("name") not in existing_names:
                ai_journals.append(fj)
        for j in (ai_journals or []):
            if isinstance(j, dict) and isinstance(j.get("match"), (int, float)):
                j["match"] = f"{j['match']}%"
        parsed_json["journals"] = ai_journals[:10]

    # ── 组装 Report ──────────────────────────────────
    sm = parsed_json.get("summary", {})

    def _safe_parse(items, cls):
        out = []
        for r in (items or []):
            if isinstance(r, dict):
                try:
                    out.append(cls(**r))
                except Exception:
                    pass
        return out

    # 解析逻辑连贯性审查
    lr_raw = parsed_json.get("logical_review", {}) or {}
    logical_review = None
    try:
        coherence_issues = []
        for ci in (lr_raw.get("coherence_issues") or []):
            if isinstance(ci, dict):
                coherence_issues.append(CoherenceIssue(
                    location=ci.get("location", "未知位置"),
                    issue_type=ci.get("issue_type", "sentence_coherence"),
                    description=ci.get("description", ""),
                    severity=ci.get("severity", "warning"),
                    suggestion=ci.get("suggestion"),
                ))
        knowledge_graph = KnowledgeGraph(
            nodes=[KnowledgeGraphNode(**n) for n in (lr_raw.get("knowledge_graph", {}).get("nodes") or [])],
            edges=[KnowledgeGraphEdge(**e) for e in (lr_raw.get("knowledge_graph", {}).get("edges") or [])],
            summary=lr_raw.get("knowledge_graph", {}).get("summary") or "",
        )
        logical_review = LogicalReview(
            research_theme=lr_raw.get("research_theme") or "",
            research_framework=lr_raw.get("research_framework") or "",
            section_logic=lr_raw.get("section_logic") or [],
            argument_logic=lr_raw.get("argument_logic") or [],
            coherence_issues=coherence_issues,
            theme_consistency=lr_raw.get("theme_consistency") or [],
            overall_assessment=lr_raw.get("overall_assessment", ""),
            knowledge_graph=knowledge_graph,
        )
    except Exception:
        pass

    # 组装报告
    report = CompletionReport(
        id=_review_id,
        file_name=parsed.file_name,
        status="completed",
        summary=ReviewSummary(
            overall_score=max(0.0, min(100.0, float(sm.get("overall_score", 60.0)))),
            strengths=list(sm.get("strengths", [])),
            weaknesses=list(sm.get("weaknesses", [])),
            recommendation=str(sm.get("recommendation", "minor_revision")),
        ),
        rules=rules,
        ai_reviews=_safe_parse(parsed_json.get("ai_reviews"), AIReviewItem),
        revisions=_safe_parse(parsed_json.get("revisions"), Revision),
        completions=_safe_parse(parsed_json.get("completions"), CompletionItem),
        logical_review=logical_review,
        polished_paper=polished_text,        # 来自并行任务
        literature_review=lit_review_text,   # 来自并行任务
        llm_success=llm_success,
        error_messages=error_messages,
    )

    # 记录历史（仅保留最近 10 篇）
    _history.append(HistoryRecord(
        id=report.id,
        file_name=parsed.file_name,
        timestamp=datetime.now(timezone.utc),
        summary={"score": report.summary.overall_score, "recommendation": report.summary.recommendation},
    ))
    _history.sort(key=lambda r: r.timestamp)
    if len(_history) > 10:
        for old in _history[:-10]:
            _reports.pop(old.id, None)
            _original_texts.pop(old.id, None)
            _original_files.pop(old.id, None)
            _journals.pop(old.id, None)
            _models_used.pop(old.id, None)
        _history[:] = _history[-10:]
    _reports[report.id] = report
    _original_texts[report.id] = text
    _journals[report.id] = ai_journals
    _models_used[report.id] = model_name or os.getenv("MODEL_NAME", "gpt-4o")

    return report


# ── 管理员认证 ──────────────────────────────────────

from jose import jwt, JWTError
import hashlib

# 如果 _ADMIN_PASSWORD 不是 SHA256 哈希（长度不为64），则进行哈希
if ADMIN_PASSWORD and len(ADMIN_PASSWORD) != 64:
    ADMIN_PASSWORD = hashlib.sha256(ADMIN_PASSWORD.encode()).hexdigest()


def _authenticate(username: str, password: str) -> bool:
    """验证管理员用户名密码"""
    import hashlib
    password_hash = hashlib.sha256(password.encode()).hexdigest()
    return username == ADMIN_USERNAME and password_hash == ADMIN_PASSWORD


def _create_token(username: str) -> str:
    """创建 JWT token"""
    from datetime import datetime, timedelta
    expire = datetime.now(timezone.utc) + timedelta(hours=JWT_EXPIRE_HOURS)
    payload = {"sub": username, "exp": expire}
    return jwt.encode(payload, JWT_SECRET, algorithm="HS256")


def _verify_token(token: str) -> dict:
    """验证 JWT token，返回 claims"""
    try:
        payload = jwt.decode(token, JWT_SECRET, algorithms=["HS256"])
        if payload.get("sub") and payload.get("exp"):
            return payload
    except JWTError:
        pass
    return {}


def _get_current_user(credentials: HTTPAuthorizationCredentials = Depends(security)) -> dict:
    """从请求中获取当前登录的管理员"""
    payload = _verify_token(credentials.credentials)
    if not payload:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="无效的认证令牌",
            headers={"WWW-Authenticate": "Bearer"},
        )
    return payload


# ── FastAPI App ────────────────────────────────────────

@asynccontextmanager
async def _lifespan(app: FastAPI):
    yield

app = FastAPI(title="论文审稿系统", version="1.0.0", lifespan=_lifespan)

app.add_middleware(
    CORSMiddleware, allow_origins=["*"], allow_methods=["*"], allow_headers=["*"],
)


@app.get("/")
def root():
    """根路径 — 返回 API 信息"""
    return {
        "service": "学术论文审稿系统",
        "version": "1.0.0",
        "api_docs": "/docs",
        "api_health": "/api/health",
        "api_models": "/api/models",
        "api_review": "/api/review (POST)",
        "frontend": "http://localhost:3000 (Vite 开发服务器)",
    }


@app.get("/api/health")
def health():
    return {"status": "ok", "service": "paper-review-system"}


@app.post("/api/login")
def login(credentials: dict = Body(...)):
    """管理员登录"""
    username = credentials.get("username", "")
    password = credentials.get("password", "")
    if not _authenticate(username, password):
        raise HTTPException(status_code=status.HTTP_401_UNAUTHORIZED, detail="用户名或密码错误")
    token = _create_token(username)
    return {"token": token, "username": username}


@app.get("/api/models")
def list_models(force: bool = False):
    """列出可用大模型（实时探测代理可用性，force=1 跳过缓存）"""
    return _get_available_models(force_refresh=force)


@app.post("/api/review")
async def review(file: UploadFile = File(...), model: str | None = Form(default=None)):
    """上传论文并执行审稿（可选指定 model）"""
    ext = Path(file.filename or "").suffix.lower()
    raw = await file.read()

    if ext in (".txt", ".md"):
        text = raw.decode("utf-8", errors="replace")
    elif ext == ".pdf":
        text = _extract_pdf_text(raw)
    elif ext == ".docx":
        text = _extract_docx_text(raw)
    else:
        return JSONResponse(status_code=400, content={"error": f"不支持的格式：{ext}。支持 .pdf / .docx / .txt / .md"})

    # 中文字数更合理（英文用单词数，中文用字符数/2 估算）
    _wc = len(text.split())
    if _wc < 200 and len(text) > 200:  # 中文文章 split 词数极少
        _wc = len(text.replace('\n', '').replace(' ', ''))
    review_id = str(uuid.uuid4())[:8]

    parsed = ParsedDocument(
        file_name=file.filename or "unknown",
        sections=extract_sections(text),
        full_text=text,
        word_count=max(_wc, 1),
    )

    # 保存原始文件，供历史下载
    _original_files[review_id] = (raw, file.filename or "unknown")

    report = await run_review(parsed, model_name=model, review_id=review_id)
    return report.model_dump()


def _require_auth(credentials: HTTPAuthorizationCredentials = Depends(security)) -> dict:
    """依赖：检查 Bearer token 认证"""
    payload = _verify_token(credentials.credentials)
    if not payload:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="无效的认证令牌",
            headers={"WWW-Authenticate": "Bearer"},
        )
    return payload


@app.get("/api/history")
def get_history(user: dict = Depends(_require_auth)):
    records = sorted(_history, key=lambda r: r.timestamp, reverse=True)
    return [r.model_dump() for r in records]


@app.delete("/api/history/{review_id}")
def delete_history(review_id: str, user: dict = Depends(_require_auth)):
    global _history
    _history = [r for r in _history if r.id != review_id]
    _reports.pop(review_id, None)
    _original_texts.pop(review_id, None)
    _original_files.pop(review_id, None)
    _journals.pop(review_id, None)
    _models_used.pop(review_id, None)
    return {"ok": True}


@app.get("/api/recommend-journals/{review_id}")
def recommend_journals_for_review(review_id: str):
    """根据审稿报告和论文内容，推荐 Top 5 投稿期刊"""
    from journal_recommender import recommend_journals as _r

    report = _reports.get(review_id)
    if report is None:
        return JSONResponse(status_code=404, content={"error": "报告不存在"})

    text = _original_texts.get(review_id, "")
    llm_journals = _journals.get(review_id, [])

    # 优先用 LLM 推荐，为空则用引擎推荐
    if llm_journals and len(llm_journals) >= 3:
        return llm_journals

    return _r(text, overall_score=float(report.summary.overall_score))


@app.get("/api/download/{review_id}")
def download_report(review_id: str):
    """下载审稿报告 DOCX（含修改痕迹和批注）"""
    from fastapi.responses import Response
    from urllib.parse import quote

    report = _reports.get(review_id)
    if report is None:
        return JSONResponse(status_code=404, content={"error": "报告不存在或已过期，请重新提交审稿"})

    # 只有 LLM 成功执行了审稿才允许下载
    if not report.llm_success:
        msgs = "; ".join(report.error_messages)
        return JSONResponse(
            status_code=403,
            content={"error": f"LLM 审稿失败，无法下载报告: {msgs}"},
        )

    original_text = _original_texts.get(review_id, "")
    journals = _journals.get(review_id)
    docx_bytes = _build_docx(report, original_text, journals)

    # 生成文件名：论文名-模型名-北京时间日期-审稿报告.docx
    safe_name = report.file_name.rsplit('.', 1)[0]
    model_name = _models_used.get(review_id, "unknown")
    model_short = model_name.rsplit('/', 1)[-1].replace(':', '-').replace('_', '-')
    bj_date = report.timestamp.astimezone(timezone(timedelta(hours=8))).strftime("%Y%m%d")
    filename = f"{safe_name}-{model_short}-{bj_date}-审稿报告.docx"
    encoded_filename = quote(filename)

    return Response(
        content=docx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"
        }
    )


@app.get("/api/download-original/{review_id}")
def download_original_file(review_id: str):
    """下载用户提交的原始稿件"""
    from fastapi.responses import Response
    from urllib.parse import quote

    data = _original_files.get(review_id)
    if data is None:
        return JSONResponse(status_code=404, content={"error": "原始稿件不存在"})
    raw_bytes, original_filename = data
    ext = Path(original_filename or "").suffix.lower()
    if ext:
        filename = f"{original_filename.rsplit('.', 1)[0]}-original{ext}"
    else:
        filename = f"{original_filename or 'unknown'}-original"
    encoded_filename = quote(filename)

    media_types = {
        ".pdf": "application/pdf",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".txt": "text/plain",
        ".md": "text/markdown",
    }
    return Response(
        content=raw_bytes,
        media_type=media_types.get(ext, "application/octet-stream"),
        headers={
            "Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"
        }
    )


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)